import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os
from settings import SETTINGS
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import subprocess
class TranscriptEmailSender:
    def __init__(self, sender_email=SETTINGS.sender_email, sender_password=SETTINGS.sender_password,
                 smtp_server: str = "smtp.gmail.com", smtp_port: int = 587):
        self.sender_email = sender_email
        self.sender_password = sender_password
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port

    def send_transcription_email(self, recipient_email: str, transcript_file_paths: list[str]):

        try:
            # Create email
            msg = MIMEMultipart()
            msg["From"] = self.sender_email
            msg["To"] = recipient_email
            msg["Subject"] = "התמלול שביקשת מוכן! מצורף בזה."

            # Add simple message body
            body = "התמלול שביקשת מוכן! מצורף בזה כקובץ."
            msg.attach(MIMEText(body, "plain"))

            # Attach all transcription files
            for file_path in transcript_file_paths:
                try:
                    with open(file_path, "rb") as f:
                        part = MIMEApplication(f.read(), Name=os.path.basename(file_path))
                        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                        msg.attach(part)
                except Exception as e:
                    print(f"⚠️ Failed to attach file {file_path}: {e}")


            # Send email
            with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                server.starttls()
                server.login(self.sender_email, self.sender_password)
                server.send_message(msg)

            print(f"✅ Transcription email sent to {recipient_email}")

        except Exception as e:
            print(f"❌ Error sending transcription email: {e}")

    def generate_transcript_pdf(self,txt_input_path: str, pdf_output_path: str):
        txt_path = Path(txt_input_path)
        pdf_path = Path(pdf_output_path)
        docx_path = pdf_path.with_suffix(".docx")

        # Read lines from input text file
        lines = txt_path.read_text(encoding="utf-8").splitlines()

        # Create DOCX document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(12)

        def set_paragraph_rtl(paragraph):
            paragraph.alignment = 2  # Right-align
            p = paragraph._element
            pPr = p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)

        # Add paragraphs with bold speaker names and RTL direction
        for line in lines:
            if ":" in line:
                speaker, text = line.split(":", 1)
                para = doc.add_paragraph()
                para.add_run(f"{speaker.strip()}: ").bold = True
                para.add_run(f"{text.strip()}")
                set_paragraph_rtl(para)

        # Save DOCX file
        doc.save(docx_path)

        # Convert DOCX to PDF using pandoc and xelatex
        subprocess.run([
            "pandoc",
            str(docx_path),
            "-o", str(pdf_path),
            "--pdf-engine=xelatex",
            "--variable", "mainfont=FreeSans",
            "--variable", "lang=he",
            "--variable", "direction=RTL",
            "--variable", "geometry=margin=2.5cm"
        ], check=True)

        print(f"✅ PDF created at: {pdf_path}")

        return 

    def run(self, recipient_email, transcript_file_path):
        
        transcript_path = Path(transcript_file_path)
        pdf_path = transcript_path.with_suffix(".pdf")
        self.generate_transcript_pdf(transcript_file_path,pdf_path)
        transcript_file_paths = [transcript_file_path, pdf_path]
        self.send_transcription_email(recipient_email, transcript_file_paths)
        